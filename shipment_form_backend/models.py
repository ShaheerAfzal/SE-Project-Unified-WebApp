# models.py
import re
from io import BytesIO
# import uuid
from django.db import models
from django.conf import settings
from django.utils import timezone
from django.core.files.base import ContentFile
from django.http import FileResponse


# Optional: if you use docxtpl for templating
# pip install docxtpl
try:
    from docxtpl import DocxTemplate
    DOCXTPL_AVAILABLE = True
except Exception:
    DOCXTPL_AVAILABLE = False
    
# Create your models here.

# helper: extract placeholders in two common formats:
def extract_placeholders_from_docx(path_or_file):
    """
    Extract placeholders from a .docx file.
    Supports:
        [KEY]
        {{KEY}}
        {KEY}
    Returns a cleaned, normalized list of keys.
    """
    import re
    from docx import Document

    # Load document from either file path or file object
    if hasattr(path_or_file, "read"):
        doc = Document(path_or_file)
    else:
        doc = Document(path_or_file)

    # Collect text
    text_blocks = []
    for p in doc.paragraphs:
        text_blocks.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_blocks.append(cell.text)

    joined = "\n".join(text_blocks)

    keys = set()

    # --- Extract [KEY] ---
    for m in re.findall(r"\[([^\]]+)\]", joined):
        keys.add(m.strip())

    # --- Extract {{KEY}} ---
    for m in re.findall(r"\{\{([^{}]+)\}\}", joined):
        keys.add(m.strip())

    # --- Extract {KEY} (your format) ---
    # Make sure NOT to capture {{KEY}} again.
    for m in re.findall(r"(?<!\{)\{([^{}]+)\}(?!\})", joined):
        keys.add(m.strip())

    # Normalize keys: Replace spaces → underscores, title-case optional
    normalized = []
    for k in keys:
        nk = re.sub(r"\s+", "_", k.strip())
        normalized.append(nk)

    return sorted(set(normalized))

#func
class DocumentTemplate(models.Model):
    """
    Stores an uploaded .docx template and its extracted field definitions.
    fields stores a mapping of placeholder_key -> friendly_label (optional).
      Example: {"PRODUCT_NAME": "Product Name", "QUANTITY": "Quantity"}
    key_field is one of the placeholder keys used as the primary identifier for documents.
    """
    # id = models.UUIDField(primary_key=True, default=uuid.uuid4, editable=False)

    name = models.CharField(max_length=255)
    file = models.FileField(upload_to="doc_templates/")
    fields = models.JSONField(default=dict, blank=True)
    key_field = models.CharField(max_length=255, blank=True, null=True)
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)

    def _str_(self):
        return self.name

    def extract_and_populate_fields(self):
        """
        Extract placeholders from the docx file and populate self.fields
        if fields is empty or if the file changed.
        Uses the extract_placeholders_from_docx helper.
        Friendly labels default to a humanized version of the key.
        """
        # open underlying file for reading
        f = self.file
        f.open(mode="rb")
        try:
            keys = extract_placeholders_from_docx(f)
        finally:
            f.close()

        # build mapping key -> friendly label
        mapping = {}
        for k in keys:
            # humanize: replace underscores with spaces, title-case
            label = k.replace("_", " ").strip().title()
            mapping[k] = label

        self.fields = mapping
        # set a sensible default key_field if not set
        if not self.key_field and keys:
            self.key_field = keys[0]
        self.save(update_fields=["fields", "key_field", "updated_at"])

    def generate_filled_docx_bytes(self, field_values: dict):
        from docx import Document
        
        # 1. Open the file safely
        if hasattr(self.file, "path"):
            doc = Document(self.file.path)
        else:
            self.file.open()
            doc = Document(self.file)
            self.file.seek(0) 

        # 2. Recursive function to replace text everywhere (paragraphs, tables, etc.)
        def replace_in_element(element, data):
            # Handle Paragraphs
            if hasattr(element, 'paragraphs'):
                for p in element.paragraphs:
                    # Optimization: Only process if it looks like a placeholder
                    if '{' in p.text or '[' in p.text:
                        for key, val in data.items():
                            val_str = str(val)
                            # Handle [KEY]
                            if f"[{key}]" in p.text:
                                p.text = p.text.replace(f"[{key}]", val_str)
                            # Handle {{KEY}}
                            if f"{{{{{key}}}}}" in p.text:
                                p.text = p.text.replace(f"{{{{{key}}}}}", val_str)
                            # Handle {KEY}
                            if f"{{{key}}}" in p.text:
                                p.text = p.text.replace(f"{{{key}}}", val_str)

            # Handle Tables (Recursive)
            if hasattr(element, 'tables'):
                for table in element.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            replace_in_element(cell, data)

        # 3. Run replacement
        replace_in_element(doc, field_values)

        # 4. Save to memory
        out = BytesIO()
        doc.save(out)
        out.seek(0)
        return out


class GeneratedDocument(models.Model):
    """
    Stores a record of a generated document (but not necessarily the binary file).
    field_values: the data used to fill the template so the doc can be regenerated.
    If you want to store the generated binary file, you can add an optional FileField.
    """
    # id = models.UUIDField(primary_key=True, default=uuid.uuid4, editable=False)
    
    template = models.ForeignKey(DocumentTemplate, on_delete=models.CASCADE, related_name="documents")
    created_by = models.ForeignKey(settings.AUTH_USER_MODEL, null=True, blank=True, on_delete=models.SET_NULL)
    created_at = models.DateTimeField(auto_now_add=True)
    field_values = models.JSONField()  # mapping key -> value
    key_field_value = models.CharField(max_length=255, blank=True, null=True)
    # optionally store the generated file (commented out)
    # file = models.FileField(upload_to="generated/", null=True, blank=True)

    def _str_(self):
        return f"{self.template.name} - {self.key_field_value or self.created_at.isoformat()}"

    def generate_and_attach_file(self):
        """
        Generate .docx bytes and optionally save to self.file if you enabled it.
        Returns a BytesIO of the generated docx.
        """
        out = self.template.generate_filled_docx_bytes(self.field_values)
        # If you want to attach/save the generated file:
        # filename = f"{self.template.name}_{self.key_field_value or timezone.now().strftime('%Y%m%d%H%M%S')}.docx"
        # self.file.save(filename, ContentFile(out.read()), save=True)
        out.seek(0)
        return out
    def save(self, *args, **kwargs):
        # Automatically set key_field_value if it's missing
        if not self.key_field_value and self.template and self.template.key_field:
            self.key_field_value = self.field_values.get(self.template.key_field)
            
        super().save(*args, **kwargs)